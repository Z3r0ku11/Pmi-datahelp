from pathlib import Path
import fitz
from docx import Document

ROOT = Path(__file__).resolve().parent.parent

FRAMEWORK_DIR = ROOT / "docs" / "frameworks"
OUTPUT_DIR = FRAMEWORK_DIR / "extracted"

OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

PDF_FILE = FRAMEWORK_DIR / "Framework Corporativo y Proceso de Gestión de Proyectos v3.1 - Final.pdf"
DOCX_FILE = FRAMEWORK_DIR / "Framework Gestión Agil de Proyectos V1.docx"

PDF_OUTPUT = OUTPUT_DIR / "Framework-Corporativo-v3.1.md"
DOCX_OUTPUT = OUTPUT_DIR / "Framework-Gestion-Agil-v1.md"


def clean_text(text):
    lines = [line.strip() for line in text.splitlines()]
    return "\n".join(line for line in lines if line)


def extract_pdf():
    print("")
    print("=" * 70)
    print("EXTRAYENDO FRAMEWORK CORPORATIVO")
    print("=" * 70)

    if not PDF_FILE.exists():
        print("ERROR: PDF no encontrado:")
        print(PDF_FILE)
        return False

    document = fitz.open(PDF_FILE)
    page_count = len(document)

    output = [
        "# Framework Corporativo y Proceso de Gestión de Proyectos v3.1",
        "",
        f"**Documento fuente:** `{PDF_FILE.name}`",
        f"**Cantidad de páginas:** {page_count}",
        "",
        "---",
        ""
    ]

    total_chars = 0
    empty_pages = 0

    for page_number, page in enumerate(document, start=1):
        text = clean_text(page.get_text("text"))

        output.append(f"# Página {page_number}")
        output.append("")
        output.append(f"<!-- SOURCE_PAGE: {page_number} -->")
        output.append("")

        if text:
            output.append(text)
            total_chars += len(text)
        else:
            output.append(
                "> [ADVERTENCIA] No se pudo extraer texto de esta página."
            )
            empty_pages += 1

        output.append("")
        output.append("---")
        output.append("")

    document.close()

    PDF_OUTPUT.write_text(
        "\n".join(output),
        encoding="utf-8"
    )

    print(f"Archivo generado: {PDF_OUTPUT}")
    print(f"Páginas PDF: {page_count}")
    print(f"Caracteres extraídos: {total_chars}")
    print(f"Páginas sin texto: {empty_pages}")

    if total_chars < 1000:
        print("")
        print("ADVERTENCIA:")
        print("El PDF contiene muy poco texto extraíble.")
        print("Podría tratarse de un PDF basado en imágenes.")

    return True


def paragraph_to_markdown(paragraph):
    text = paragraph.text.strip()

    if not text:
        return None

    try:
        style = paragraph.style.name.lower()
    except Exception:
        style = ""

    if "title" in style:
        return f"# {text}"

    if "heading 1" in style:
        return f"# {text}"

    if "heading 2" in style:
        return f"## {text}"

    if "heading 3" in style:
        return f"### {text}"

    if "heading 4" in style:
        return f"#### {text}"

    if "list" in style:
        return f"- {text}"

    return text


def extract_docx():
    print("")
    print("=" * 70)
    print("EXTRAYENDO FRAMEWORK ÁGIL")
    print("=" * 70)

    if not DOCX_FILE.exists():
        print("ERROR: DOCX no encontrado:")
        print(DOCX_FILE)
        return False

    document = Document(DOCX_FILE)

    output = [
        "# Framework Gestión Ágil de Proyectos V1",
        "",
        f"**Documento fuente:** `{DOCX_FILE.name}`",
        "",
        "---",
        ""
    ]

    paragraph_count = 0

    for paragraph in document.paragraphs:
        markdown = paragraph_to_markdown(paragraph)

        if markdown:
            output.append(markdown)
            output.append("")
            paragraph_count += 1

    if document.tables:
        output.append("# Tablas del documento")
        output.append("")

    for table_number, table in enumerate(document.tables, start=1):
        output.append(f"## Tabla {table_number}")
        output.append("")

        rows = []

        for row in table.rows:
            cells = []

            for cell in row.cells:
                value = (
                    cell.text
                    .strip()
                    .replace("\n", " ")
                    .replace("|", "\\|")
                )
                cells.append(value)

            rows.append(cells)

        if not rows:
            continue

        column_count = max(len(row) for row in rows)

        normalized = [
            row + [""] * (column_count - len(row))
            for row in rows
        ]

        output.append(
            "| " + " | ".join(normalized[0]) + " |"
        )

        output.append(
            "| " + " | ".join(["---"] * column_count) + " |"
        )

        for row in normalized[1:]:
            output.append(
                "| " + " | ".join(row) + " |"
            )

        output.append("")

    DOCX_OUTPUT.write_text(
        "\n".join(output),
        encoding="utf-8"
    )

    print(f"Archivo generado: {DOCX_OUTPUT}")
    print(f"Párrafos extraídos: {paragraph_count}")
    print(f"Tablas extraídas: {len(document.tables)}")

    return True


def main():
    print("")
    print("=" * 70)
    print("PMO FRAMEWORK HUB - EXTRACCIÓN DOCUMENTAL")
    print("=" * 70)

    pdf_ok = extract_pdf()
    docx_ok = extract_docx()

    print("")
    print("=" * 70)

    if pdf_ok and docx_ok:
        print("EXTRACCIÓN FINALIZADA CORRECTAMENTE")
    else:
        print("EXTRACCIÓN FINALIZADA CON ERRORES")

    print("=" * 70)
    print("")


if __name__ == "__main__":
    main()